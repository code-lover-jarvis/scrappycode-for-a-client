import urllib.request, urllib.parse, urllib.error
from bs4 import BeautifulSoup
import docx
import re
import os
import pandas as pd

mydoc = docx.Document()
mydoc.core_properties.author = 'Jibin Joy'
try: os.mkdir('wikiimages')
except:
    print('Please remove the already existing wikiimages folder. Exiting....')
    quit()
print('''\n\n\n\n\n\n\n\n>>>This is a project made by Jibin Joy, \nin which the user could enter a wikipedia URL and \nsave the page as a docx file.
\nHope this project helps students to create assignments<<<\n\n\n\n\n\n''')

url = input('Enter the URL : ')
print('\n\n\n')
count = 0
html = urllib.request.urlopen(url).read()
soup = BeautifulSoup(html, 'html.parser')
html = soup.find("div", {"id": "content"})
for i in html.find_all():
    if(i.name == 'h1'):
        mydoc.add_heading(i.get_text(), 0)
        p = mydoc.add_paragraph('\n\n')
    if(i.name == 'h2'):
        mydoc.add_page_break()
        mydoc.add_heading(i.get_text(), 1)
        p = mydoc.add_paragraph('\n\n')



    if(i.name == 'table'):
        try:
            if(i['class'][0] == 'wikitable'):
                for j in i.find_all():
                    if(j.name == 'th' or j.name == 'td'):
                        p.add_run(j.get_text().strip() + '\t')

        except: pass



    if(i.name == 'img'):
        try:
            if(i['class'][0] == 'thumbimage'):
                count+=1
                images = i['src']
                urllib.request.urlretrieve('https:'+images, 'wikiimages\imgname{}.jpg'.format(count))
                mydoc.add_picture('wikiimages\imgname{}.jpg'.format(count))
                p = mydoc.add_paragraph('\n\n')

        except:
            print("\n\n \t\t\t:(\nEncountered an error in adding an image to the document...\n\nPlease wait...\n\n")



    if(i.name == 'p'):
        txt = i.get_text()
        txt = re.sub("\[[0-9]+\]", "", txt)
        txt = re.sub("\[[A-Z]\]", "", txt)
        mydoc.add_paragraph(txt)
        p = mydoc.add_paragraph('\n\n')





print("\t\tDon't worry!!! We have saved the image files in the folder wikiimages")



mydoc.save(r'wikifile.docx')
print('''Done!!!!\n Check for a 'wikifile.docx' \n\n Thanks for making use of my project.''')
